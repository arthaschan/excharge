#!/usr/bin/env python3
"""把三份老师点评答复文档合并成一份 docx，便于打印/转发。"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = [
    'docs/老师点评_逐条答复记录.md',
    'docs/给老师回应_可解释性实现说明.md',
    'docs/期刊升级方案_可解释性反哺检测.md',
]
OUT = 'docs/TAIG摘要点评答复_合并版.docx'

doc = Document()
# 中文字体
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_runs(p, text):
    """解析 **bold** 与 `code` 内联格式。"""
    for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Courier New'
        else:
            p.add_run(part)


def add_table(doc, lines):
    rows = [[c.strip() for c in ln.strip().strip('|').split('|')] for ln in lines]
    data = [r for r in rows if not all(re.match(r'^:?-{2,}:?$', c) for c in r)]
    if not data:
        return
    ncol = len(data[0])
    t = doc.add_table(rows=len(data), cols=ncol)
    t.style = 'Table Grid'
    for ri, row in enumerate(data):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs(p, row[ci] if ci < len(row) else '')
            if ri == 0:
                for r in p.runs:
                    r.bold = True


def process(path, top_level=1):
    lines = open(path, encoding='utf-8').read().split('\n')
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1; continue
        if s.startswith('|'):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            add_table(doc, tbl); continue
        if s.startswith('#### '):
            doc.add_heading(s[5:], level=4)
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        elif s.startswith('## '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('# '):
            doc.add_heading(s[2:], level=top_level)
        elif s == '---':
            pass
        elif s.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = add_runs(p, s[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, s[2:])
        elif re.match(r'^\d+[.、]\s*', s):
            p = doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\d+[.、]\s*', '', s))
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
        i += 1


# 封面标题
title = doc.add_heading('TAIG 摘要点评答复（合并版）', level=0)
intro = doc.add_paragraph()
add_runs(intro, '本文档合并以下三份材料，用于向评审老师说明摘要的逐条答复、可解释性实现细节与期刊升级方向。')
intro2 = doc.add_paragraph()
add_runs(intro2, '① 逐条答复记录　② 可解释性实现说明　③ 期刊升级方案')
intro2.paragraph_format.space_after = Pt(18)

# 三份文档（第一份用 level=1，后两份用 level=1 分节）
for idx, path in enumerate(SRC):
    process(path, top_level=1)
    if idx < len(SRC) - 1:
        doc.add_page_break()

doc.save(OUT)
print('Saved', OUT)
